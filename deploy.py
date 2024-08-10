import typer as t #pip install typer
from art import text2art #pip install art
from termcolor import colored #pip install termcolor
import os
from pathlib import Path
from docx import Document #pip install python-docx
import docx2txt   #pip install docx2txt


def honeytoken_deploy():

    while True:
        p_input=input("Specify the path (absolute path): ")
        path=Path(p_input)

        if p_input == "exit":
            break

        if not path.is_absolute():
            t.echo(colored("\n\t\tPlease provide the absolute path\n","yellow"))
        else:
            break
    
    if not path.parent.exists():
        try:
            path.mkdir(parents=True,exist_ok=True)
            t.echo("\nCreated directories for the path: " + colored(path,"magenta") + "\n")
        except Exception as e:
            t.echo(colored(f"An error occured while creating directories: {e}","yellow"))
    elif path.exists():
        t.echo("\nCurrent path: " + colored(path,"magenta") + "\n")
    else:
        path.mkdir(exist_ok=True)
        t.echo(f"Created directories for the path: {path}")

    t.echo(colored("\t\tSupported file types","green"))
    t.echo(colored("\t\t-[","green") + colored(".txt ","magenta")+ colored("]-  Text File","green"))
    t.echo(colored("\t\t-[","green") + colored(".conf","magenta")+ colored("]-  Configuration File","green"))
    t.echo(colored("\t\t-[","green") + colored(".ini ","magenta")+ colored("]-  Initialization File","green"))
    t.echo(colored("\t\t-[","green") + colored(".docx","magenta")+ colored("]-  Microsoft Word Document (Open XML Document)","green"))
    t.echo(colored("\t\t-[","green") + colored(".sql ","magenta")+ colored("]-  Structured Query Language File","green"))
    t.echo(colored("\t\t-[","green") + colored(".pem ","magenta")+ colored("]-  Privacy Enhanced Mail File","green"))
    t.echo(colored("\t\t-[","green") + colored(".env ","magenta")+ colored("]-  Environment File","green"))
    t.echo(colored("\t\t-[","green") + colored(".json","magenta")+ colored("]-  JavaScript Object Notation File","green"))
    t.echo(colored("\t\t-[","green") + colored(".log ","magenta")+ colored("]-  Log File","green"))

    while True:
        f_name=input("\nFile name (with extension): ")
        ext=Path(f_name).suffix
        if ext in [".txt",".conf",".ini",".sql",".log",".pem",".env",".json"]:
            f_create(f_name,path)
            break
        elif ext == ".docx":
            docx(f_name,path)
            break
        else:
            t.echo(colored("\n\t\tUnsupported file format\n","yellow"))
        


def f_create(f_name,path):
    ext=Path(f_name).suffix
    default=""" """
    t.echo(colored("\n\t\t-[","green") + colored("-d","magenta") + colored("]- for default content"))
    t.echo(colored("\t\t-[","green") + colored("-c","magenta") + colored("]- for custom content\n"))
    if ext == ".txt":                   #For text file
        default="""# Sensitive Information

Username: admin
Password: P@ssw0rd1234

API Key: 12345-abcde-67890-fghij
Secret Token: 98765-zyxwv-43210-lmnop
"""
    elif ext == ".conf":                #For Configuration File
        default="""# Application Configuration

database {
    host = "db.example.com"
    port = 5432
    user = "dbuser"
    password = "dbpass"
}

server {
    host = "0.0.0.0"
    port = 8080
}

logging {
    level = "INFO"
    file = "/var/log/application.log"
}
"""
    elif ext == ".ini":                  #For Initialization File
        default="""[Database]
host = db.example.com
port = 5432
username = dbuser
password = dbpass

[Server]
host = 0.0.0.0
port = 8080

[Logging]
logfile = /var/log/app.log
loglevel = INFO
"""
    elif ext == ".sql":                     #For Structured Query Language File
        default="""-- Database Schema

CREATE TABLE users (
    id SERIAL PRIMARY KEY,
    username VARCHAR(50) UNIQUE NOT NULL,
    password VARCHAR(255) NOT NULL,
    email VARCHAR(100) UNIQUE
);

CREATE TABLE logs (
    id SERIAL PRIMARY KEY,
    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    level VARCHAR(20),
    message TEXT
);

INSERT INTO users (username, password, email) VALUES
('admin', 'P@ssw0rd1234', 'admin@example.com'),
('user', 'User1234', 'user@example.com');
"""
    elif ext == ".log":                         #For log file
        default="""[2024-08-07 10:00:00] ERROR - Failed to connect to database: db.example.com:5432
[2024-08-07 10:05:00] INFO - Server started on 0.0.0.0:8080
[2024-08-07 10:10:00] WARN - Deprecated API used in request: GET /api/v1/old-endpoint
[2024-08-07 10:15:00] ERROR - Unexpected token in JSON response from API
"""
    elif ext == ".pem":                    #For Privacy Enhanced Mail File
        default="""-----BEGIN CERTIFICATE-----
MIIDXTCCAkWgAwIBAgIJAKP4KVIePu1jMA0GCSqGSIb3DQEBCwUAMEUxCzAJBgNV
BAYTAkFVMSUwIwYDVQQKExxJbnRlcm5ldCBXaWRnaXRzIFB0eSBMdGQuMRcwFQYD
VQQDEw5teXNpdGUuY29tLmF1MB4XDTIxMDkxMDEyMjk1NloXDTIxMTAxMDEyMjk1
NlowRTELMAkGA1UEBhMCQVUxJTAjBgNVBAoTHEludGVybmV0IFdpZGdpdHMgUHR5
IEx0ZC4xFzAVBgNVBAMTDm15c2l0ZS5jb20uYXUwggEiMA0GCSqGSIb3DQEBAQUA
A4IBDwAwggEKAoIBAQC9jTY3QHxu6jCqJl/1KrVv6xL6p/kNfnVGZjWY8C3kXUFA
ry7X2ZKJ/vlgEcv37V4wDwKsV6GtnH0vPd1l1yRQWMoW4ICgKyn5yAj7h2fZcEGH
LfDP4Bq9mv+vPcy1jEIJW9qAK9wFyQPoCmf7hTyzRPw4UP9z3eYfT6bT+F6lSAtf
AGnDd27uLFe/UX8PBKEmHUpEvhq6B+VkDImll1MeYX/fQvY3LQZxTVVsaSRJ18tb
LZ6+OzLp/klC5l0mI+zSi+brarO/uyjqykvcmkkM4QpNc4QIlgGTcTkcA7mCeSKj
/gzKbQXe1Jb4Am4BAkB8OsU9ZY5h3GyK8jlmXv0TAgMBAAGjUDBOMB0GA1UdDgQW
BBTm5nrGf5/3ylLCs2L96Kv8g6MaJTAfBgNVHSMEGDAWgBTm5nrGf5/3ylLCs2L9
6Kv8g6MaJTAMBgNVHRMEBTADAQH/MA0GCSqGSIb3DQEBCwUAA4IBAQBCfA74J4PT
rSbOvE8Rk00h5aX3W36kP5UKn6X5vv+7TkTmM3mHg3Zg0Tj6GiZ9Jg4/P+Gb1U+e
nkAr6xnPgg1JFXZqwq45lvhKXl6ZWUtN3sYeC2CCKfLZjdyNJnSxJ1h/jwllBy+j
WDB0+eDH9+DpO/6aaeb4qU5kzFeH0dIv2yIDTYEYHZDWoSmNfGmTwhHC6Y+VmD32
Qb6bTq2X9eC1vYbh9/cP0YeXl5P4OqZwqE5jMBH7pJX5HCx0Gxv1r/pycUikXyRa
1kBlSTyWKNmrm7DF0NQl4sTnMmSwQ0sFZMHXt6BjBn5KO9QDQ9PXlfk+yKGoXvOs
KdjDxeONF8Zk
-----END CERTIFICATE-----
"""
    elif ext == ".env":                    #For Environment File
        default="""# Environment Variables

DATABASE_URL=postgres://dbuser:dbpass@db.example.com:5432/mydatabase
SECRET_KEY=supersecretkey12345
DEBUG=True
API_KEY=12345-abcde-67890-fghij
"""
    elif ext ==".json":                   #For JavaScript Object Notation File
        default="""{
    "database": {
        "host": "db.example.com",
        "port": 5432,
        "username": "dbuser",
        "password": "dbpass"
    },
    "server": {
        "host": "0.0.0.0",
        "port": 8080
    },
    "logging": {
        "level": "INFO",
        "file": "/var/log/application.log"
    }
}
"""

    try:                                       #To create the file
        with open(path/f_name,"w+") as file:
            while True:
                ct=input("Select one: ")
                if ct == "-d":
                    file.write(default)
                    t.echo(colored("\n\tFile created\nFile content:\n","blue"))
                    file.seek(0)
                    c=file.read()
                    t.echo(c)
                    break
                elif ct == "-c":
                    file.write(custom())
                    t.echo(colored("\n\tFile created\nFile content:\n","blue"))
                    file.seek(0)
                    c=file.read()
                    t.echo(c)
                    t.echo("\n")
                    break
                else:
                    t.echo(colored("\n\t\tInvalid input\n","yellow"))
    except Exception as e:
        t.echo(colored("An error occurred: ","yellow") + colored(e,"red"))
        


def docx(f_name,path):
    doc=Document()
    t.echo(colored("\n\t\t-[","green") + colored("-d","magenta") + colored("]- for default content"))
    t.echo(colored("\t\t-[","green") + colored("-c","magenta") + colored("]- for custom content\n"))
    while True:
        ct=input("Select one: ")
        if ct == "-d":
            doc.add_heading('Confidential Report', level=1)
            con="""
Date: August 7, 2024

Project Name: New Initiative

Executive Summary:
The project aims to revolutionize our approach to data management, enhancing security and efficiency.

Key Findings:
- Significant improvements in data processing times.
- Enhanced security protocols.

Recommendations:
1. Implement updated encryption methods.
2. Conduct a thorough audit of current practices.

Contact Information:
John Doe
Chief Technology Officer
Email: john.doe@example.com
"""
            doc.add_paragraph(con)
            break
        elif ct=="-c":
            t.echo(colored("\n\t\tGive the heading (In one line)","blue"))
            head=input()
            doc.add_heading(head,level=1)
            doc.add_paragraph(custom())
            break
        else:
            t.echo(colored("\n\t\tInvalid input\n","yellow"))
    try:
        doc_path=path/f_name
        doc.save(doc_path)
        t.echo(colored("\n\tFile created\nFile content:\n","blue"))
        txt=docx2txt.process(doc_path)
        t.echo(txt,"\n")
    except Exception as e:
        t.echo(colored("An error occurred: ","yellow") + colored(e,"red"))
            


def custom():                        #To get the customized input from the user to add in the file
    t.echo(colored("\n\tEnter your custom content (type '-END' on a new line to finish)\n","blue"))
    lines=[]
    while True:
        l=input()
        if l == "-END" or l=="-end":
            break
        lines.append(l)
    
    return '\n'.join(lines)